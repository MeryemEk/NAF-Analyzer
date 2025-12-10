"""
French Sector Analyzer - Streamlit App
======================================

A comprehensive tool for analyzing French business sectors using INSEE data,
BOAMP public tenders, and geographic analysis.

Author: REKOLT Analysis
"""

import streamlit as st
import pandas as pd
import numpy as np
import io
import tempfile
import shutil
from pathlib import Path
import requests
import traceback

# Page configuration
st.set_page_config(
    page_title="French Sector Analyzer",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
    <style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #071F3B;
        margin-bottom: 1rem;
    }
    .sub-header {
        font-size: 1.5rem;
        color: #10433E;
        margin-top: 2rem;
        margin-bottom: 1rem;
    }
    .info-box {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
    .success-box {
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
    .warning-box {
        background-color: #fff3cd;
        border: 1px solid #ffeaa7;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
    </style>
""", unsafe_allow_html=True)

# Initialize session state
if 'analysis_complete' not in st.session_state:
    st.session_state.analysis_complete = False
if 'results' not in st.session_state:
    st.session_state.results = {}
if 'temp_dir' not in st.session_state:
    st.session_state.temp_dir = None


def download_boamp_data(keywords: str, progress_callback=None) -> pd.DataFrame:
    """
    Download BOAMP data based on keywords query.
    
    Parameters
    ----------
    keywords : str
        Search keywords for BOAMP query
    progress_callback : callable, optional
        Function to call with progress updates
        
    Returns
    -------
    pd.DataFrame
        BOAMP data
    """
    if progress_callback:
        progress_callback("Downloading BOAMP data...")
    
    BASE_URL = "https://boamp-datadila.opendatasoft.com/api/records/1.0/download/"
    
    params = {
        "dataset": "boamp",
        "q": keywords,
        "refine.etat": "INITIAL",
        "format": "csv",
        "csv_separator": ";",
    }
    
    try:
        resp = requests.get(BASE_URL, params=params, timeout=60)
        resp.raise_for_status()
        
        # Save to temporary file and load
        df = pd.read_csv(io.StringIO(resp.content.decode('utf-8')), delimiter=";")
        
        if progress_callback:
            progress_callback(f"Downloaded {len(df)} BOAMP records")
        
        return df
    except Exception as e:
        raise Exception(f"Error downloading BOAMP data: {str(e)}")


def validate_naf_file(df: pd.DataFrame) -> tuple[bool, str]:
    """
    Validate that uploaded NAF file has required columns.
    
    Parameters
    ----------
    df : pd.DataFrame
        Uploaded dataframe
        
    Returns
    -------
    tuple[bool, str]
        (is_valid, error_message)
    """
    required_columns = [
        'siren', 'siret', 
        'activitePrincipaleEtablissement',
        'activitePrincipaleUniteLegale',
        'etatAdministratifUniteLegale',
        'etatAdministratifEtablissement',
        'codeCommuneEtablissement'
    ]
    
    missing_columns = [col for col in required_columns if col not in df.columns]
    
    if missing_columns:
        return False, f"Missing required columns: {', '.join(missing_columns)}"
    
    return True, ""


def run_sector_analysis(
    naf_file_path: str,
    naf_code: str,
    sector_name: str,
    boamp_keywords: str,
    pop_filepath: str,
    output_dir: str,
    progress_bar,
    status_text
) -> dict:
    """
    Run complete sector analysis pipeline.
    
    Parameters
    ----------
    naf_file_path : str
        Path to NAF CSV file
    naf_code : str
        NAF code (e.g., '43.13Z')
    sector_name : str
        Name of the sector for reports
    boamp_keywords : str
        Keywords for BOAMP query
    pop_filepath : str
        Path to population data file
    output_dir : str
        Directory for output files
    progress_bar : streamlit.progress
        Progress bar object
    status_text : streamlit.empty
        Status text object
        
    Returns
    -------
    dict
        Results from all analyses
    """
    results = {}
    
    try:
        # Import analysis modules
        status_text.text("Loading analysis modules...")
        progress_bar.progress(0.05)
        
        from sector_entities_analyzer import run_analysis as run_entities_analysis
        from sector_offer_analyzer import run_analysis as run_offer_analysis
        from sector_demand_analyzer import run_analysis as run_demand_analysis
        from sector_opportunity_analyzer import run_opportunity_analysis
        
        # 1. Download BOAMP data
        status_text.text("Downloading BOAMP public tender data...")
        progress_bar.progress(0.10)
        
        def boamp_progress(msg):
            status_text.text(f"BOAMP: {msg}")
        
        boamp_df = download_boamp_data(boamp_keywords, boamp_progress)
        boamp_path = Path(output_dir) / "boamp_data.csv"
        boamp_df.to_csv(boamp_path, sep=";", index=False)
        progress_bar.progress(0.20)
        
        # 2. Entities Analysis
        status_text.text("Running entities analysis...")
        progress_bar.progress(0.25)
        
        entities_results = run_entities_analysis(
            filepath=naf_file_path,
            naf_code=naf_code,
            output_dir=str(Path(output_dir) / "entities")
        )
        results['entities'] = entities_results
        progress_bar.progress(0.40)
        
        # 3. Offer Analysis
        status_text.text("Running geographic offer analysis...")
        progress_bar.progress(0.45)
        
        offer_results = run_offer_analysis(
            filepath=naf_file_path,
            naf_code=naf_code,
            output_dir=str(Path(output_dir) / "offer"),
            pop_filepath=pop_filepath
        )
        results['offer'] = offer_results
        progress_bar.progress(0.60)
        
        # 4. Demand Analysis
        status_text.text("Running demand analysis...")
        progress_bar.progress(0.65)
        
        demand_results = run_demand_analysis(
            boamp_filepath=str(boamp_path),
            output_dir=str(Path(output_dir) / "demand"),
            pop_filepath=pop_filepath,
            second_demand_filepath=None,  # Not using second demand source
            second_demand_name=None,
            second_demand_source=None
        )
        results['demand'] = demand_results
        progress_bar.progress(0.80)
        
        # 5. Opportunity Analysis
        status_text.text("Computing market opportunities...")
        progress_bar.progress(0.85)
        
        opportunity_results = run_opportunity_analysis(
            df_demand_clusters=demand_results['demand_clusters_combined'],
            df_offer_clusters=offer_results['offre_clustered'],
            departements_grouped=offer_results['departements_grouped'],
            total_per_dept=demand_results['boamp_total_per_dept'],
            output_dir=output_dir
    
        )
        results['opportunities'] = opportunity_results
        progress_bar.progress(0.95)
        
        # Store metadata
        results['metadata'] = {
            'naf_code': naf_code,
            'sector_name': sector_name,
            'boamp_keywords': boamp_keywords,
            'boamp_records': len(boamp_df)
        }
        
        status_text.text("Analysis complete!")
        progress_bar.progress(1.0)
        
        return results
        
    except Exception as e:
        status_text.text(f"Error: {str(e)}")
        st.error(f"Analysis failed: {str(e)}")
        st.error(traceback.format_exc())
        raise


def main():
    """Main Streamlit app."""
    
    # Header
    st.markdown('<div class="main-header">French Sector Analyzer</div>', unsafe_allow_html=True)
    st.markdown("**Comprehensive analysis of French business sectors using INSEE data and public tenders**")
    
    # Sidebar for inputs
    st.sidebar.header("Analysis Parameters")
    
    # File uploads
    st.sidebar.subheader("1. Upload NAF Data")
    naf_file = st.sidebar.file_uploader(
        "Upload INSEE NAF CSV file. You can download it from https://annuaire-entreprises.data.gouv.fr/export-sirene",
        type=['csv'],
        help="Upload the INSEE company registry CSV file for your sector"
    )
    
    # Sector parameters
    st.sidebar.subheader("2. Sector Information")
    naf_code = st.sidebar.text_input(
        "NAF Code",
        placeholder="e.g., 43.13Z or 81.29A",
        help="Enter the NAF code for this sector"
    )
    
    sector_name = st.sidebar.text_input(
        "Sector Name",
        placeholder="e.g., Drilling and boring",
        help="Enter a descriptive name for this sector"
    )
    
    # BOAMP keywords
    st.sidebar.subheader("3. BOAMP Search Keywords")
    boamp_keywords = st.sidebar.text_area(
        "Keywords Query",
        placeholder='("keyword1" OR "keyword2" OR "phrase with spaces")',
        help="Enter search keywords for BOAMP public tenders. Use OR for multiple terms.",
        height=100
    )
    
    # Instructions
    with st.sidebar.expander("Instructions & Data Format"):
        st.markdown("""
        **Required NAF File Columns:**
        - siren, siret
        - activitePrincipaleEtablissement
        - activitePrincipaleUniteLegale
        - etatAdministratifUniteLegale
        - etatAdministratifEtablissement
        - codeCommuneEtablissement
        - dateCreationEtablissement
        - dateCreationUniteLegale
        
        **BOAMP Keywords Format:**
        Use boolean operators:
        - OR: `("term1" OR "term2")`
        - AND: `"term1" AND "term2"`
        - Phrases: `"exact phrase"`
        
        **Example:**
        ```
        ("drilling" OR "boring" OR "soil investigation")
        ```
        """)
    
    # Run analysis button
    run_analysis_btn = st.sidebar.button(
        "Run Analysis",
        type="primary",
        use_container_width=True,
        disabled=not (naf_file and naf_code and sector_name and boamp_keywords)
    )
    
    # Main content area
    if not naf_file:
        st.info("Please upload an INSEE NAF CSV file to begin analysis.")
        st.markdown("""
        <div class="info-box">
        <h3>Welcome to French Sector Analyzer</h3>
        <p>This tool performs comprehensive analysis of French business sectors by combining:</p>
        <ul>
            <li><strong>Entity Analysis:</strong> Evolution of companies and establishments</li>
            <li><strong>Offer Analysis:</strong> Geographic distribution of supply</li>
            <li><strong>Demand Analysis:</strong> Public tender activity (BOAMP)</li>
            <li><strong>Opportunity Analysis:</strong> Market gaps and opportunities</li>
        </ul>
        <p>To get started, upload your data and fill in the parameters in the sidebar.</p>
        </div>
        """, unsafe_allow_html=True)
        return
    
    # Validate NAF file
    if naf_file:
        try:
            naf_df = pd.read_csv(naf_file, dtype=str, low_memory=False)
            is_valid, error_msg = validate_naf_file(naf_df)
            
            if not is_valid:
                st.error(f"Invalid NAF file: {error_msg}")
                return
            
            st.success(f"NAF file loaded: {len(naf_df):,} records, {naf_df['siren'].nunique():,} unique legal entities")
            
        except Exception as e:
            st.error(f"Error reading NAF file: {str(e)}")
            return
    
    # Run analysis
    if run_analysis_btn:
        st.markdown('<div class="sub-header">Running Analysis</div>', unsafe_allow_html=True)
        
        # Create temporary directory
        temp_dir = tempfile.mkdtemp()
        st.session_state.temp_dir = temp_dir
        
        try:
            # Save uploaded file
            naf_path = Path(temp_dir) / "naf_data.csv"
            naf_df.to_csv(naf_path, index=False)
            
            # Copy required data files to temp directory
            pop_path = Path(temp_dir) / "popdept.xlsx"
            # We'll need to include these files with the app
            
            # Progress tracking
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            # Run analysis
            results = run_sector_analysis(
                naf_file_path=str(naf_path),
                naf_code=naf_code,
                sector_name=sector_name,
                boamp_keywords=boamp_keywords,
                pop_filepath="popdept.xlsx",  # Will be in app directory
                output_dir=temp_dir,
                progress_bar=progress_bar,
                status_text=status_text
            )
            
            # Store results in session state
            st.session_state.results = results
            st.session_state.analysis_complete = True
            
            st.success("Analysis completed successfully!")
            
        except Exception as e:
            st.error(f"Analysis failed: {str(e)}")
            st.error("Please check your input files and parameters.")
            return
    
    # Display results
    if st.session_state.analysis_complete and st.session_state.results:
        display_results(st.session_state.results)


def display_results(results: dict):
    """
    Display analysis results in the app.
    
    Parameters
    ----------
    results : dict
        Results from sector analysis
    """
    st.markdown('<div class="sub-header">Analysis Results</div>', unsafe_allow_html=True)
    
    metadata = results.get('metadata', {})
    
    # Summary metrics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric(
            "Sector",
            metadata.get('sector_name', 'N/A')
        )
    
    with col2:
        st.metric(
            "NAF Code",
            metadata.get('naf_code', 'N/A')
        )
    
    with col3:
        entities = results.get('entities', {})
        if 'data' in entities:
            st.metric(
                "Legal Entities",
                f"{entities['data']['siren'].nunique():,}"
            )
    
    with col4:
        st.metric(
            "BOAMP Records",
            f"{metadata.get('boamp_records', 0):,}"
        )
    
    # Tabs for different analyses
    tabs = st.tabs([
        "Entity Analysis",
        "Offer Analysis",
        "Demand Analysis",
        "Opportunities",
        "Generate Report"
    ])
    
    # Entity Analysis Tab
    with tabs[0]:
        display_entity_analysis(results.get('entities', {}))
    
    # Offer Analysis Tab
    with tabs[1]:
        display_offer_analysis(results.get('offer', {}))
    
    # Demand Analysis Tab
    with tabs[2]:
        display_demand_analysis(results.get('demand', {}))
    
    # Opportunities Tab
    with tabs[3]:
        display_opportunity_analysis(results.get('opportunities'))
    
    # Generate Report Tab
    with tabs[4]:
        display_report_generation(results)


def display_entity_analysis(entities_results: dict):
    """Display entity analysis results."""
    st.subheader("Entity Analysis")
    st.markdown("Analysis of company creation, evolution, and survival rates.")
    
    if not entities_results:
        st.warning("No entity analysis results available.")
        return
    
    # Key metrics
    col1, col2, col3 = st.columns(3)
    
    data = entities_results.get('data')
    if data is not None:
        with col1:
            st.metric(
                "Total Establishments",
                f"{len(data):,}"
            )
        
        with col2:
            active_legal = data[data['etatAdministratifUniteLegale'] == 'A']
            st.metric(
                "Active Legal Entities",
                f"{active_legal['siren'].nunique():,}"
            )
        
        with col3:
            active_estab = data[data['etatAdministratifEtablissement'] == 'A']
            st.metric(
                "Active Establishments",
                f"{len(active_estab):,}"
            )
    
    # Charts section
    st.markdown("---")
    
    # Look for saved charts in output directory
    if st.session_state.temp_dir:
        output_dir = Path(st.session_state.temp_dir) / "entities"
        
        chart_files = [
            ("01_units_evolution.png", "Evolution of Legal Units"),
            ("02_establishments_evolution.png", "Evolution of Establishments"),
            ("03_distribution.png", "Distribution of Establishments per Entity"),
            ("04_top10_entities.png", "Top 10 Entities"),
            ("06_greenfield_brownfield.png", "Greenfield vs Brownfield"),
            ("07_survival_by_duration.png", "Survival Rates by Duration"),
        ]
        
        for filename, title in chart_files:
            chart_path = output_dir / filename
            if chart_path.exists():
                st.subheader(title)
                st.image(str(chart_path), use_container_width=True)
            
    # Show data tables
    st.markdown("---")
    st.subheader("Data Tables")
    
    if 'units_per_year' in entities_results:
        with st.expander("Units Created Per Year"):
            st.dataframe(entities_results['units_per_year'], use_container_width=True)
    
    if 'establishments_per_year' in entities_results:
        with st.expander("Establishments Created Per Year"):
            st.dataframe(entities_results['establishments_per_year'], use_container_width=True)
    
    if 'top10' in entities_results:
        with st.expander("Top 10 Entities by Number of Establishments"):
            st.dataframe(entities_results['top10'], use_container_width=True)


def display_offer_analysis(offer_results: dict):
    """Display offer analysis results."""
    st.subheader("Geographic Offer Analysis")
    st.markdown("Distribution of establishments across French departments.")
    
    if not offer_results:
        st.warning("No offer analysis results available.")
        return
    
    offre = offer_results.get('offre_clustered')
    if offre is not None:
        # Summary stats
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric(
                "Average per 100k",
                f"{offre['establishments_per_100k'].mean():.1f}"
            )
        
        with col2:
            st.metric(
                "Median per 100k",
                f"{offre['establishments_per_100k'].median():.1f}"
            )
        
        with col3:
            high_supply = offre[offre['cluster_offre'] == 'high supply']
            st.metric(
                "High Supply Depts",
                f"{len(high_supply)}"
            )
        
        # Display maps and charts
        st.markdown("---")
        
        if st.session_state.temp_dir:
            output_dir = Path(st.session_state.temp_dir) / "offer"
            
            chart_files = [
                ("01_offer_map.png", "Establishments per 100k Inhabitants by Department"),
                ("02_offer_histogram.png", "Distribution of Establishment Density"),
                ("03_offer_clusters.png", "Supply Clusters by Department"),
            ]
            
            for filename, title in chart_files:
                chart_path = output_dir / filename
                if chart_path.exists():
                    st.subheader(title)
                    st.image(str(chart_path), use_container_width=True)
        
        # Data table
        st.markdown("---")
        st.subheader("Establishments per Department")
        
        display_df = offre[['dept_group', 'dept_name', 'establishments_count', 
                           'establishments_per_100k', 'cluster_offre']].sort_values(
            'establishments_per_100k', ascending=False
        )
        
        st.dataframe(display_df, use_container_width=True)


def display_demand_analysis(demand_results: dict):
    """Display demand analysis results."""
    st.subheader("Demand Analysis (BOAMP Public Tenders)")
    st.markdown("Analysis of public tender activity across departments.")
    
    if not demand_results:
        st.warning("No demand analysis results available.")
        return
    
    # BOAMP summary
    boamp_per_year = demand_results.get('boamp_per_year')
    if boamp_per_year is not None:
        st.subheader("BOAMP Activity Over Time")
        
        col1, col2 = st.columns(2)
        with col1:
            st.dataframe(boamp_per_year, use_container_width=True)
        
        with col2:
            if len(boamp_per_year) > 0:
                latest_year = boamp_per_year['Year'].max()
                latest_count = boamp_per_year[boamp_per_year['Year'] == latest_year]['Number of Unique BOAMP Publications'].values[0]
                st.metric(f"Tenders in {int(latest_year)}", f"{int(latest_count):,}")
    
    # Display charts
    st.markdown("---")
    
    if st.session_state.temp_dir:
        output_dir = Path(st.session_state.temp_dir) / "demand"
        
        # Look for saved PNG files
        import os
        if output_dir.exists():
            png_files = sorted([f for f in os.listdir(output_dir) if f.endswith('.png')])
            
            for png_file in png_files:
                chart_path = output_dir / png_file
                # Create a title from filename
                title = png_file.replace('.png', '').replace('_', ' ').title()
                st.subheader(title)
                st.image(str(chart_path), use_container_width=True)
    
    # Geographic distribution
    st.markdown("---")
    total_per_dept = demand_results.get('boamp_total_per_dept')
    if total_per_dept is not None:
        st.subheader("BOAMP Tenders by Department")
        
        
        # Check which columns exist
        available_cols = ['dept_group']
        if 'dept_name' in total_per_dept.columns:
            available_cols.append('dept_name')
        if 'total_idweb' in total_per_dept.columns:
            available_cols.append('total_idweb')
        if 'idweb_per_100k' in total_per_dept.columns:
            available_cols.append('idweb_per_100k')
        
        display_df = total_per_dept[available_cols].sort_values(
             available_cols[-1], ascending=False
        )

        st.dataframe(display_df, use_container_width=True)
    
    # Demand clusters
    demand_clusters = demand_results.get('demand_clusters_combined')
    if demand_clusters is not None:
        st.markdown("---")
        st.subheader("Demand Clusters")
        
        cluster_summary = demand_clusters['cluster_demand'].value_counts()
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Active Colonies", cluster_summary.get('active colonies', 0))
        with col2:
            st.metric("Breeding Grounds", cluster_summary.get('breeding grounds', 0))
        with col3:
            st.metric("Dormant Nests", cluster_summary.get('dormant nests', 0))
        
        with st.expander("View Demand Clusters Data"):
            st.dataframe(demand_clusters, use_container_width=True)


def display_opportunity_analysis(opportunities):
    """Display opportunity analysis results."""
    st.subheader("Market Opportunity Analysis")
    st.markdown("Confrontation of supply and demand to identify market opportunities.")
    
    if opportunities is None:
        st.warning("No opportunity analysis results available.")
        return
    
    # Opportunity distribution
    if isinstance(opportunities, pd.DataFrame):
        st.subheader("Opportunity Matrix")
        
        opp_summary = opportunities['opportunity'].value_counts()
        
        # Display metrics in two rows
        cols1 = st.columns(3)
        cols2 = st.columns(3)
        
        metrics_to_show = [
            ('white space', 'White Space', cols1[0]),
            ('growth frontiers', 'Growth Frontiers', cols1[1]),
            ('crowded spots', 'Crowded Spots', cols1[2]),
            ('emerging competition', 'Emerging Competition', cols2[0]),
            ('quiet zones', 'Quiet Zones', cols2[1]),
            ('over-served areas', 'Over-served Areas', cols2[2]),
        ]
        
        for opp_key, opp_label, col in metrics_to_show:
            with col:
                st.metric(opp_label, opp_summary.get(opp_key, 0))
        
        # Display opportunity map
        st.markdown("---")
        
        if st.session_state.temp_dir:
            # Look for opportunity map in the output directory
            possible_paths = [
                Path(st.session_state.temp_dir) / "opportunity_map.png",
                Path(st.session_state.temp_dir) / "opportunities" / "opportunity_map.png",
            ]
            
            for chart_path in possible_paths:
                if chart_path.exists():
                    st.subheader("Market Opportunity Map")
                    st.image(str(chart_path), use_container_width=True)
                    break
        
        # Full table
        st.markdown("---")
        st.subheader("Opportunity Details by Department")
        
        # Sort by priority
        display_df = opportunities.sort_values('priority')
        st.dataframe(display_df, use_container_width=True)
        
        # Opportunity definitions
        with st.expander("Opportunity Category Definitions"):
            st.markdown("""
            **Priority 1 - White Space**: High demand, low supply
            - Departments with active tender activity but limited competition
            - Primary market opportunity for expansion
            
            **Priority 2 - Growth Frontiers**: Medium demand, low supply
            - Emerging markets with growth potential
            - Secondary market opportunity
            
            **Priority 3 - Crowded Spots**: High demand, high supply
            - Competitive markets with active tender activity
            - Requires differentiation strategy
            
            **Priority 4 - Emerging Competition**: Medium demand, high supply
            - Markets with increasing competition
            - Monitor for market changes
            
            **Priority 5 - Quiet Zones**: Low demand, low supply
            - Limited market activity
            - Lower priority for expansion
            
            **Priority 6 - Over-served Areas**: Low demand, high supply
            - Saturated markets with limited growth
            - Lowest priority for new entry
            """)


def display_report_generation(results: dict):
    """Display report generation options."""
    st.subheader("Generate Analysis Report")
    st.markdown("Create a comprehensive Word document report with all analysis results.")
    
    st.markdown("""
    <div class="info-box">
    <h4>Report Contents:</h4>
    <ul>
        <li>Executive summary with key metrics</li>
        <li>Entity analysis charts and tables</li>
        <li>Geographic offer analysis with maps</li>
        <li>Demand analysis from BOAMP data</li>
        <li>Market opportunity matrix and maps</li>
        <li>Commentary sections for your insights</li>
    </ul>
    </div>
    """, unsafe_allow_html=True)
    
    if st.button("Generate Word Report", type="primary"):
        with st.spinner("Generating report..."):
            try:
                report_path = generate_word_report(results)
                
                with open(report_path, 'rb') as f:
                    st.download_button(
                        label="Download Report",
                        data=f,
                        file_name=f"sector_analysis_{results['metadata']['naf_code'].replace('.', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                st.success("Report generated successfully!")
                
            except Exception as e:
                st.error(f"Error generating report: {str(e)}")


def generate_word_report(results: dict) -> str:
    """
    Generate Word document report with analysis results.
    
    Parameters
    ----------
    results : dict
        Analysis results
        
    Returns
    -------
    str
        Path to generated report file
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import os
    
    doc = Document()
    
    # Set document properties
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # Add title
    title = doc.add_heading(f"Sector Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading(f"{results['metadata']['sector_name']}", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata section
    doc.add_paragraph()
    metadata_para = doc.add_paragraph()
    metadata_para.add_run(f"NAF Code: ").bold = True
    metadata_para.add_run(f"{results['metadata']['naf_code']}\n")
    metadata_para.add_run(f"Analysis Date: ").bold = True
    metadata_para.add_run(f"{pd.Timestamp.now().strftime('%B %d, %Y')}\n")
    metadata_para.add_run(f"BOAMP Records Analyzed: ").bold = True
    metadata_para.add_run(f"{results['metadata']['boamp_records']:,}\n")
    metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add page break
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading("Executive Summary", 1)
    
    # Key metrics table
    entities = results.get('entities', {})
    if 'data' in entities:
        data = entities['data']
        active_entities = data[data['etatAdministratifUniteLegale'] == 'A']
        active_estab = data[data['etatAdministratifEtablissement'] == 'A']
        
        doc.add_paragraph(f"Total Legal Entities: {data['siren'].nunique():,}")
        doc.add_paragraph(f"Active Legal Entities: {active_entities['siren'].nunique():,}")
        doc.add_paragraph(f"Total Establishments: {len(data):,}")
        doc.add_paragraph(f"Active Establishments: {len(active_estab):,}")
    
    doc.add_paragraph()
    doc.add_paragraph("[COMMENTARY SECTION]")
    doc.add_paragraph("Please add your executive summary and key takeaways here.")
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # 1. Entity Analysis
    doc.add_heading("1. Entity Analysis", 1)
    doc.add_paragraph("Analysis of company creation, evolution, and survival rates in the sector.")
    doc.add_paragraph()
    
    # Add entity analysis charts
    if st.session_state.temp_dir:
        output_dir = Path(st.session_state.temp_dir) / "entities"
        
        chart_files = [
            ("01_units_evolution.png", "Evolution of Legal Units by Creation Year"),
            ("02_establishments_evolution.png", "Evolution of Establishments by Creation Year"),
            ("03_distribution.png", "Distribution of Establishments per Entity"),
            ("04_top10_entities.png", "Top 10 Entities by Number of Establishments"),
            ("06_greenfield_brownfield.png", "Greenfield vs Brownfield Establishments"),
            ("07_survival_by_duration.png", "Survival Rates by Duration of Existence"),
        ]
        
        for filename, caption in chart_files:
            chart_path = output_dir / filename
            if chart_path.exists():
                doc.add_heading(caption, level=2)
                try:
                    doc.add_picture(str(chart_path), width=Inches(6))
                except Exception as e:
                    doc.add_paragraph(f"[Chart: {caption} - Image could not be loaded]")
                doc.add_paragraph()
                doc.add_paragraph("[COMMENTARY SECTION]")
                doc.add_paragraph(f"Please add your analysis of the {caption.lower()} here.")
                doc.add_paragraph()
    
    doc.add_page_break()
    
    # 2. Offer Analysis
    doc.add_heading("2. Geographic Offer Analysis", 1)
    doc.add_paragraph("Distribution of establishments across French departments.")
    doc.add_paragraph()
    
    offer = results.get('offer', {})
    if 'offre_clustered' in offer:
        offre_df = offer['offre_clustered']
        doc.add_paragraph(f"Average establishments per 100k inhabitants: {offre_df['establishments_per_100k'].mean():.1f}")
        doc.add_paragraph(f"Median establishments per 100k inhabitants: {offre_df['establishments_per_100k'].median():.1f}")
        
        high_supply = offre_df[offre_df['cluster_offre'] == 'high supply']
        doc.add_paragraph(f"High supply departments: {len(high_supply)}")
        doc.add_paragraph()
    
    # Add offer analysis charts
    if st.session_state.temp_dir:
        output_dir = Path(st.session_state.temp_dir) / "offer"
        
        chart_files = [
            ("01_offer_map.png", "Establishments per 100k Inhabitants by Department"),
            ("02_offer_histogram.png", "Distribution of Establishment Density"),
            ("03_offer_clusters.png", "Supply Clusters by Department"),
        ]
        
        for filename, caption in chart_files:
            chart_path = output_dir / filename
            if chart_path.exists():
                doc.add_heading(caption, level=2)
                try:
                    doc.add_picture(str(chart_path), width=Inches(6))
                except Exception as e:
                    doc.add_paragraph(f"[Chart: {caption} - Image could not be loaded]")
                doc.add_paragraph()
                doc.add_paragraph("[COMMENTARY SECTION]")
                doc.add_paragraph(f"Please add your analysis of the {caption.lower()} here.")
                doc.add_paragraph()
    
    doc.add_page_break()
    
    # 3. Demand Analysis
    doc.add_heading("3. Demand Analysis (BOAMP Public Tenders)", 1)
    doc.add_paragraph("Analysis of public tender activity across departments.")
    doc.add_paragraph()
    
    demand = results.get('demand', {})
    
    if 'boamp_per_year' in demand:
        boamp_df = demand['boamp_per_year']
        latest_year = boamp_df['Year'].max()
        latest_count = boamp_df[boamp_df['Year'] == latest_year]['Number of Unique BOAMP Publications'].values[0]
        doc.add_paragraph(f"BOAMP tenders in {int(latest_year)}: {int(latest_count):,}")
        doc.add_paragraph()
    
    # Add demand analysis charts
    if st.session_state.temp_dir:
        output_dir = Path(st.session_state.temp_dir) / "demand"
        
        if output_dir.exists():
            png_files = sorted([f for f in os.listdir(output_dir) if f.endswith('.png')])
            
            for png_file in png_files:
                chart_path = output_dir / png_file
                caption = png_file.replace('.png', '').replace('_', ' ').title()
                
                doc.add_heading(caption, level=2)
                try:
                    doc.add_picture(str(chart_path), width=Inches(6))
                except Exception as e:
                    doc.add_paragraph(f"[Chart: {caption} - Image could not be loaded]")
                doc.add_paragraph()
                doc.add_paragraph("[COMMENTARY SECTION]")
                doc.add_paragraph(f"Please add your analysis here.")
                doc.add_paragraph()
    
    doc.add_page_break()
    
    # 4. Opportunity Analysis
    doc.add_heading("4. Market Opportunity Analysis", 1)
    doc.add_paragraph("Confrontation of supply and demand to identify market opportunities.")
    doc.add_paragraph()
    
    opportunities = results.get('opportunities')
    if opportunities is not None and isinstance(opportunities, pd.DataFrame):
        opp_summary = opportunities['opportunity'].value_counts()
        
        doc.add_paragraph("Opportunity distribution across departments:")
        for opp_type, count in opp_summary.items():
            doc.add_paragraph(f"  • {opp_type.title()}: {count} departments")
        doc.add_paragraph()
    
    # Add opportunity map
    if st.session_state.temp_dir:
        possible_paths = [
            Path(st.session_state.temp_dir) / "opportunity_map.png",
            Path(st.session_state.temp_dir) / "opportunities" / "opportunity_map.png",
        ]
        
        for chart_path in possible_paths:
            if chart_path.exists():
                doc.add_heading("Market Opportunity Map", level=2)
                try:
                    doc.add_picture(str(chart_path), width=Inches(6))
                except Exception as e:
                    doc.add_paragraph("[Chart: Market Opportunity Map - Image could not be loaded]")
                doc.add_paragraph()
                break
    
    doc.add_paragraph("[COMMENTARY SECTION]")
    doc.add_paragraph("Please add your strategic analysis of market opportunities here.")
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # 5. Conclusions and Recommendations
    doc.add_heading("5. Strategic Conclusions and Recommendations", 1)
    
    doc.add_paragraph("[COMMENTARY SECTION]")
    doc.add_paragraph()
    doc.add_paragraph("Please add your overall conclusions and strategic recommendations here, including:")
    doc.add_paragraph("  • Key market insights")
    doc.add_paragraph("  • Priority departments for expansion")
    doc.add_paragraph("  • Competitive positioning recommendations")
    doc.add_paragraph("  • Risk factors and considerations")
    doc.add_paragraph("  • Next steps and action items")
    
    # Save report
    report_path = Path(st.session_state.temp_dir) / "sector_analysis_report.docx"
    doc.save(report_path)
    
    return str(report_path)


if __name__ == "__main__":
    main()
